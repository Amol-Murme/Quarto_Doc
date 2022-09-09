import docx
doc = docx.Document('hello.docx')

len(doc.paragraphs)

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

text = getText("test.docx")   
# print(text)

###########################################
# for paragraph in doc.paragraphs:
    # print(1)
    # for run in paragraph.runs:
    #     print(run._element.xml)
    #     if 'lastRenderedPageBreak' in run._element.xml:
    #         print ('soft page break found at run:', run.text[:20]) 
    #     if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
    #         print ('hard page break found at run:', run.text[:20])
    # count = 0 
    # for run in paragraph.runs:
    #     print(run._element.br_lst)
    #     count +=1
    #     if run._element.br_lst:             
    #         for br in run._element.br_lst:
    #             # br_couter+=1
    #             print (br.type)  

    # print(count)    


#############################################

level_from_style_name = {f'Heading {i}': i for i in range(10)}

def format_levels(cur_lev):
    levs = [str(l) for l in cur_lev if l != 0]
    return '.'.join(levs)  # Customize your format here

d = docx.Document('test.docx')

current_levels = [0] * 10
full_text = []

for p in d.paragraphs:
    if p.style.name not in level_from_style_name:
        full_text.append(p.text)
    else:
        level = level_from_style_name[p.style.name]
        current_levels[level] += 1
        for l in range(level + 1, 10):
            current_levels[l] = 0
        full_text.append(format_levels(current_levels) + ' ' + p.text)

# for l in full_text:
    # print(l)

#################################################################
import xml.etree.ElementTree as tree

NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
pages = []
PARA = NAMESPACE + 'p'
TEXT = NAMESPACE + 't'
PAGE = NAMESPACE + 'lastRenderedPageBreak'

aggText = ''
for paragraph in tree.iter(PARA):
     aggText += ''.join([node.text
             for node in paragraph.iter(TEXT)
             if node.text])
     if aggText and [node for node in paragraph.iter(PAGE)]:
        pages.append(aggText)
        aggText = ''
if aggText != '':
     pages.append(aggText)

print(pages)     